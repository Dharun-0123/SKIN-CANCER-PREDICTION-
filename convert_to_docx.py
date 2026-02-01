"""
Convert PROJECT_REPORT.md to a professionally formatted DOCX file.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_page_margins(doc):
    """Set page margins: Left 1.5in, Right 1in, Top 1in, Bottom 1in"""
    for section in doc.sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

def create_styles(doc):
    """Create custom styles for the document"""
    styles = doc.styles
    
    # Title Style
    if 'ProjectTitle' not in [s.name for s in styles]:
        title_style = styles.add_style('ProjectTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

def add_title_page(doc):
    """Add the title page"""
    # Main Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SKINCARE AI: AN AI-POWERED SKIN LESION CLASSIFICATION SYSTEM USING DEEP LEARNING")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A PROJECT REPORT")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted in partial fulfillment of the requirements for the award of the degree of")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MASTER OF COMPUTER APPLICATIONS (MCA)")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted by")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Student Name]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Registration Number]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Under the guidance of")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Guide Name]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Designation]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DEPARTMENT OF COMPUTER APPLICATIONS")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[INSTITUTION NAME]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[UNIVERSITY NAME]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[CITY, STATE]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ACADEMIC YEAR 2024-2025")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_page_break()


def add_certificate(doc):
    """Add certificate page"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CERTIFICATE")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    text = '''This is to certify that the project report entitled "SKINCARE AI: AN AI-POWERED SKIN LESION CLASSIFICATION SYSTEM USING DEEP LEARNING" submitted by [Student Name] bearing Registration Number [Registration Number] in partial fulfillment of the requirements for the award of the degree of Master of Computer Applications is a bonafide record of work carried out under my guidance and supervision.

This project report has not been submitted to any other university or institution for the award of any degree or diploma.'''
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("Date:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run("Place:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("[Guide Name]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("[Designation]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run("Department of Computer Applications")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("[Head of Department Name]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("Head of Department")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run("Department of Computer Applications")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_page_break()


def add_declaration(doc):
    """Add declaration page"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DECLARATION")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    text = '''I hereby declare that the project report entitled "SKINCARE AI: AN AI-POWERED SKIN LESION CLASSIFICATION SYSTEM USING DEEP LEARNING" submitted to [University Name] in partial fulfillment of the requirements for the award of the degree of Master of Computer Applications is a record of original work done by me under the guidance of [Guide Name], [Designation], Department of Computer Applications, [Institution Name].

I further declare that this project report has not been submitted to any other university or institution for the award of any degree or diploma.'''
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("Date:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run("Place:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("[Student Name]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("[Registration Number]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_page_break()


def add_acknowledgement(doc):
    """Add acknowledgement page"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ACKNOWLEDGEMENT")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    ack_text = """I would like to express my sincere gratitude to all those who have contributed to the successful completion of this project work.

First and foremost, I express my heartfelt thanks to the Almighty for blessing me with the strength, wisdom, and perseverance to complete this project successfully.

I am deeply indebted to my project guide, [Guide Name], [Designation], Department of Computer Applications, for the valuable guidance, constant encouragement, and constructive suggestions throughout the course of this project. The technical expertise and professional approach demonstrated by my guide have been instrumental in shaping this project.

I extend my sincere thanks to [Head of Department Name], Head of the Department of Computer Applications, for providing the necessary facilities and support for the successful completion of this project.

I am grateful to [Principal Name], Principal, [Institution Name], for providing an excellent academic environment and infrastructure that facilitated the completion of this project.

I would like to thank all the faculty members of the Department of Computer Applications for their valuable suggestions and support during the project development phase.

I express my gratitude to the non-teaching staff of the department for their cooperation and assistance in various administrative matters.

I am thankful to my classmates and friends for their moral support, encouragement, and valuable discussions that helped me overcome various challenges during the project development.

Finally, I express my deepest gratitude to my parents and family members for their unconditional love, support, and encouragement throughout my academic journey. Their sacrifices and blessings have been the driving force behind all my achievements."""
    
    for para in ack_text.split('\n\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(para.strip())
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("[Student Name]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    
    doc.add_page_break()


def add_abstract(doc):
    """Add abstract page"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ABSTRACT")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    abstract_text = """Skin cancer represents one of the most prevalent forms of cancer worldwide, with millions of new cases diagnosed annually. Early detection of skin cancer significantly improves treatment outcomes and survival rates, with melanoma survival rates exceeding ninety-nine percent when detected at early stages compared to merely twenty-seven percent for late-stage diagnoses. However, access to dermatological expertise remains limited in many regions, creating a critical need for accessible preliminary screening tools.

This project presents SkinCare AI, a comprehensive web-based skin lesion classification system that leverages advanced deep learning techniques to analyze dermoscopic images and provide preliminary assessments of skin conditions. The system employs a dual-model architecture combining EfficientNetB0, a state-of-the-art convolutional neural network pre-trained on ImageNet and fine-tuned on the ISIC 2019 dataset comprising 25,331 dermoscopic images, with a custom Convolutional Neural Network trained on a modified HAM10000 dataset containing 5,906 images.

The proposed system classifies skin lesions into eight distinct categories: Melanocytic Nevi, Melanoma, Basal Cell Carcinoma, Benign Keratosis-like Lesions, Actinic Keratoses, Squamous Cell Carcinoma, Vascular Lesions, and Dermatofibroma. The intelligent model selection mechanism automatically chooses the most appropriate model based on confidence thresholds, ensuring robust predictions across diverse image qualities and conditions.

The web application is developed using the Django framework with Python as the primary programming language, TensorFlow and Keras for deep learning implementation, and SQLite for database management. The system incorporates comprehensive features including secure user authentication with OTP-based email verification, analysis history tracking, PDF report generation, analytics dashboard with interactive visualizations, and an AI-powered chatbot named DermaGenie for skin health guidance.

Experimental results demonstrate that the custom CNN model achieves a test accuracy of 94.1 percent on the secondary dataset, while the EfficientNetB0 model achieves 71.32 percent accuracy on the larger ISIC 2019 dataset. The dual-model approach provides a balance between accuracy and robustness, with the system automatically falling back to the secondary model when primary model confidence is below the threshold.

The system emphasizes responsible AI usage in healthcare contexts by incorporating legally-compliant result presentation, clear medical disclaimers, and recommendations for professional medical consultation. SkinCare AI serves as an educational and preliminary screening tool, encouraging users to seek professional dermatological evaluation when potential concerns are identified."""
    
    for para in abstract_text.split('\n\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(para.strip())
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run = p.add_run("Skin Cancer Detection, Deep Learning, Convolutional Neural Networks, EfficientNetB0, Transfer Learning, Medical Image Analysis, Django, TensorFlow, Healthcare AI")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_page_break()


def add_toc(doc):
    """Add Table of Contents placeholder"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TABLE OF CONTENTS")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    toc_entries = [
        ("", "Certificate", "ii"),
        ("", "Declaration", "iii"),
        ("", "Acknowledgement", "iv"),
        ("", "Abstract", "v"),
        ("", "Table of Contents", "vi"),
        ("", "List of Tables", "ix"),
        ("", "List of Figures", "x"),
        ("1", "INTRODUCTION", "1"),
        ("1.1", "Introduction Overview", "1"),
        ("1.2", "Salient Features of the System", "4"),
        ("1.3", "Project Motivation", "7"),
        ("1.4", "Scope of the Project", "10"),
        ("1.5", "Organization of the Report", "13"),
        ("2", "SYSTEM STUDY AND ANALYSIS", "15"),
        ("2.1", "Problem Statement", "15"),
        ("2.2", "Existing System", "18"),
        ("2.2.1", "Drawbacks of Existing System", "21"),
        ("2.3", "Proposed System", "24"),
        ("2.3.1", "Advantages of Proposed System", "27"),
        ("2.4", "Feasibility Analysis", "30"),
        ("2.4.1", "Technical Feasibility", "30"),
        ("2.4.2", "Economic Feasibility", "33"),
        ("2.4.3", "Operational Feasibility", "35"),
        ("3", "DEVELOPMENT ENVIRONMENT", "38"),
        ("3.1", "Hardware Requirements", "38"),
        ("3.2", "Software Requirements", "41"),
        ("3.3", "Software Description", "44"),
        ("4", "SYSTEM DESIGN", "63"),
        ("4.1", "Module Description", "63"),
        ("4.2", "Methodology", "87"),
        ("4.3", "Input Design", "91"),
        ("4.4", "Output Design", "94"),
        ("4.5", "Data Flow Diagram", "97"),
        ("4.6", "Architecture Diagram", "102"),
        ("4.7", "Database Design", "107"),
        ("5", "SYSTEM IMPLEMENTATION", "116"),
        ("5.1", "Frontend Implementation", "116"),
        ("5.2", "Backend Implementation", "120"),
        ("5.3", "Machine Learning Model Implementation", "124"),
        ("5.4", "Model Training Process", "128"),
        ("5.5", "Model Evaluation Metrics", "132"),
        ("5.6", "Security Implementation", "136"),
        ("6", "SYSTEM TESTING", "140"),
        ("6.1", "Testing Strategy", "140"),
        ("6.2", "Unit Testing", "143"),
        ("6.3", "Integration Testing", "146"),
        ("6.4", "Validation Testing", "149"),
        ("6.5", "Test Case Design", "152"),
        ("6.6", "Sample Test Cases", "155"),
        ("7", "RESULTS AND DISCUSSION", "160"),
        ("7.1", "Experimental Results", "160"),
        ("7.2", "Performance Analysis", "164"),
        ("7.3", "Accuracy Analysis", "168"),
        ("7.4", "Comparison with Existing Systems", "172"),
        ("8", "CONCLUSION AND FUTURE ENHANCEMENT", "176"),
        ("8.1", "Conclusion", "176"),
        ("8.2", "Future Enhancements", "179"),
        ("9", "BIBLIOGRAPHY AND REFERENCES", "183"),
        ("10", "APPENDICES", "187"),
    ]
    
    # Create TOC table
    table = doc.add_table(rows=len(toc_entries), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, (num, title, page) in enumerate(toc_entries):
        row = table.rows[i]
        
        cell0 = row.cells[0]
        cell0.text = num
        for para in cell0.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if num and '.' not in num:
                    run.font.bold = True
        
        cell1 = row.cells[1]
        cell1.text = title
        for para in cell1.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if num and '.' not in num:
                    run.font.bold = True
        
        cell2 = row.cells[2]
        cell2.text = page
        for para in cell2.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    doc.add_page_break()


def add_list_of_tables(doc):
    """Add List of Tables"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LIST OF TABLES")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    tables_list = [
        ("2.1", "Comparison of Existing and Proposed System", "29"),
        ("3.1", "Minimum Hardware Requirements", "39"),
        ("3.2", "Recommended Hardware Requirements", "40"),
        ("3.3", "Software Requirements", "42"),
        ("3.4", "Python Libraries and Dependencies", "47"),
        ("4.1", "User Table Structure", "112"),
        ("4.2", "UserProfile Table Structure", "113"),
        ("4.3", "UserPredictModel Table Structure", "114"),
        ("4.4", "EmailOTP Table Structure", "115"),
        ("5.1", "EfficientNetB0 Model Specifications", "125"),
        ("5.2", "Custom CNN Model Specifications", "126"),
        ("5.3", "ISIC 2019 Dataset Distribution", "129"),
        ("5.4", "HAM10000 Subset Distribution", "130"),
        ("6.1", "User Authentication Test Cases", "156"),
        ("6.2", "Image Analysis Test Cases", "157"),
        ("6.3", "Profile Management Test Cases", "158"),
        ("6.4", "System Integration Test Cases", "159"),
        ("7.1", "Model Performance Comparison", "165"),
        ("7.2", "Class-wise Accuracy Analysis", "169"),
        ("7.3", "Comparison with Related Works", "173"),
    ]
    
    table = doc.add_table(rows=len(tables_list), cols=3)
    for i, (num, title, page) in enumerate(tables_list):
        row = table.rows[i]
        row.cells[0].text = f"Table {num}"
        row.cells[1].text = title
        row.cells[2].text = page
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()

def add_list_of_figures(doc):
    """Add List of Figures"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LIST OF FIGURES")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    figures_list = [
        ("4.1", "System Architecture Diagram", "103"),
        ("4.2", "Data Flow Diagram Level 0", "98"),
        ("4.3", "Data Flow Diagram Level 1", "100"),
        ("4.4", "Entity Relationship Diagram", "108"),
        ("4.5", "ML Workflow Diagram", "105"),
        ("4.6", "Deployment Architecture Diagram", "106"),
        ("5.1", "EfficientNetB0 Architecture", "127"),
        ("5.2", "Custom CNN Architecture", "127"),
        ("5.3", "Training Loss Curve", "131"),
        ("5.4", "Training Accuracy Curve", "131"),
        ("7.1", "Confusion Matrix - EfficientNetB0", "166"),
        ("7.2", "Confusion Matrix - Custom CNN", "167"),
        ("7.3", "ROC Curve Analysis", "170"),
        ("7.4", "Precision-Recall Curve", "171"),
        ("A.1", "Landing Page", "187"),
        ("A.2", "User Registration Page", "188"),
        ("A.3", "Login Page", "189"),
        ("A.4", "Home Dashboard", "190"),
        ("A.5", "Image Analysis Page", "191"),
        ("A.6", "Analysis Results Page", "192"),
        ("A.7", "User Profile Page", "193"),
        ("A.8", "Analytics Dashboard", "194"),
    ]
    
    table = doc.add_table(rows=len(figures_list), cols=3)
    for i, (num, title, page) in enumerate(figures_list):
        row = table.rows[i]
        row.cells[0].text = f"Figure {num}"
        row.cells[1].text = title
        row.cells[2].text = page
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()


def add_chapter_heading(doc, chapter_num, title):
    """Add a chapter heading"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"CHAPTER {chapter_num}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.upper())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()

def add_section_heading(doc, section_num, title):
    """Add a section heading"""
    p = doc.add_paragraph()
    run = p.add_run(f"{section_num} {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)

def add_subsection_heading(doc, section_num, title):
    """Add a subsection heading"""
    p = doc.add_paragraph()
    run = p.add_run(f"{section_num} {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def add_paragraph_text(doc, text):
    """Add justified paragraph text"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Inches(0.5)

def add_diagram_placeholder(doc, figure_num, title):
    """Add a diagram placeholder"""
    doc.add_paragraph()
    
    # Add placeholder box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[DIAGRAM PLACEHOLDER]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Insert diagram image here]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.italic = True
    
    doc.add_paragraph()
    
    # Add caption
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figure {figure_num}: {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    
    doc.add_paragraph()


def add_chapter1(doc):
    """Add Chapter 1: Introduction"""
    add_chapter_heading(doc, "1", "INTRODUCTION")
    
    add_section_heading(doc, "1.1", "Introduction Overview")
    
    intro_paras = [
        "The field of medical imaging has witnessed remarkable advancements in recent years, particularly with the integration of artificial intelligence and deep learning technologies. Among the various applications of these technologies in healthcare, skin cancer detection has emerged as a critical area where automated systems can significantly impact patient outcomes. Skin cancer, characterized by the abnormal growth of skin cells, represents one of the most common forms of cancer globally, affecting millions of individuals each year. The World Health Organization estimates that between two and three million non-melanoma skin cancers and approximately 132,000 melanoma skin cancers occur globally each year. These statistics underscore the urgent need for effective screening and early detection mechanisms.",
        
        "The human skin, being the largest organ of the body, serves as the primary barrier against environmental hazards and plays a crucial role in maintaining overall health. However, prolonged exposure to ultraviolet radiation from the sun, genetic predisposition, and various environmental factors can lead to the development of abnormal skin growths that may potentially become cancerous. The early identification of these abnormal growths is paramount for successful treatment and improved patient outcomes.",
        
        "Traditional methods of skin cancer diagnosis rely heavily on visual examination by trained dermatologists, followed by dermoscopic analysis and histopathological examination of biopsied tissue samples. While these methods remain the gold standard for diagnosis, they present several challenges including limited availability of dermatological expertise in rural and underserved areas, subjective interpretation that may vary among practitioners, and the time-consuming nature of the diagnostic process. These limitations have created a pressing need for automated screening tools that can assist in the preliminary assessment of skin lesions.",
        
        "The advent of deep learning, particularly Convolutional Neural Networks, has revolutionized the field of medical image analysis. These sophisticated algorithms can learn complex patterns and features from large datasets of medical images, enabling them to identify subtle characteristics that may be indicative of malignancy. Research has demonstrated that well-trained deep learning models can achieve diagnostic accuracy comparable to, and in some cases exceeding, that of experienced dermatologists.",
        
        "SkinCare AI represents a comprehensive solution that harnesses the power of deep learning to provide accessible, preliminary skin lesion screening. The system is designed as a web-based application that allows users to upload images of skin lesions and receive instant AI-powered analysis. By combining state-of-the-art neural network architectures with a user-friendly interface, SkinCare AI aims to bridge the gap between advanced medical technology and everyday users who may benefit from early screening.",
        
        "The system employs a dual-model architecture that combines the strengths of two distinct neural network models. The primary model utilizes EfficientNetB0, a highly efficient convolutional neural network architecture that has demonstrated excellent performance on image classification tasks while maintaining computational efficiency. This model has been pre-trained on the ImageNet dataset and subsequently fine-tuned on the ISIC 2019 dataset, which comprises over 25,000 dermoscopic images representing various skin conditions.",
        
        "The secondary model is a custom-designed Convolutional Neural Network that has been trained on a modified version of the HAM10000 dataset. This dataset, originally published by Tschandl and colleagues in 2018, contains dermoscopic images of common pigmented skin lesions and has become a benchmark dataset for skin lesion classification research. The custom CNN model provides an alternative classification pathway that can be utilized when the primary model exhibits low confidence in its predictions.",
        
        "The intelligent model selection mechanism implemented in SkinCare AI automatically evaluates the confidence of predictions from the primary model and determines whether to utilize the secondary model for improved accuracy. This approach ensures that users receive the most reliable predictions possible, regardless of variations in image quality or lesion characteristics.",
        
        "Beyond the core classification functionality, SkinCare AI incorporates a comprehensive suite of features designed to enhance the user experience and provide valuable insights. These features include secure user authentication with email verification, detailed analysis history tracking, professional PDF report generation for medical consultations, an interactive analytics dashboard for visualizing trends, and an AI-powered chatbot named DermaGenie that provides guidance on skin health topics.",
        
        "The development of SkinCare AI has been guided by principles of responsible AI usage in healthcare contexts. The system incorporates clear medical disclaimers, emphasizes that results are for educational purposes only, and consistently encourages users to seek professional medical evaluation for any concerns. This approach ensures that the technology serves as a complement to, rather than a replacement for, professional dermatological care."
    ]
    
    for para in intro_paras:
        add_paragraph_text(doc, para)

    
    add_section_heading(doc, "1.2", "Salient Features of the System")
    
    features_paras = [
        "SkinCare AI incorporates a comprehensive array of features that distinguish it from conventional skin analysis applications and establish it as a sophisticated healthcare technology solution. The system has been designed with careful consideration of user needs, technical requirements, and the sensitive nature of medical information processing. This section provides a detailed examination of the key features that define the SkinCare AI platform.",
        
        "The dual-model architecture represents one of the most significant technical innovations implemented in SkinCare AI. Unlike traditional single-model systems that may struggle with certain image types or conditions, the dual-model approach provides redundancy and improved accuracy across a wider range of scenarios. The primary model, based on the EfficientNetB0 architecture, has been trained on the extensive ISIC 2019 dataset comprising over 25,000 dermoscopic images. This model excels at identifying subtle patterns and features in high-quality dermoscopic images. The secondary model, a custom-designed Convolutional Neural Network trained on a modified HAM10000 dataset, provides an alternative classification pathway that can be particularly effective for images that may not conform to typical dermoscopic standards.",
        
        "The intelligent model selection mechanism automatically evaluates the confidence level of predictions from the primary model and determines whether to utilize the secondary model for potentially improved accuracy. This approach ensures that users receive the most reliable predictions possible, regardless of variations in image quality, lighting conditions, or lesion characteristics. The system can operate in three modes: automatic selection, where the system chooses the most appropriate model based on confidence thresholds; EfficientNetB0-only mode, where users specifically request analysis using the primary model; and CNN-only mode, where users prefer the secondary model for their analysis.",
        
        "The classification capability of SkinCare AI extends to eight distinct categories of skin conditions, providing comprehensive coverage of the most common skin lesions encountered in clinical practice. These categories include Melanocytic Nevi, which are commonly known as moles and represent benign proliferations of melanocytes; Melanoma, the most dangerous form of skin cancer that develops from melanocytes; Basal Cell Carcinoma, the most common type of skin cancer that rarely metastasizes but can cause significant local tissue destruction; Benign Keratosis-like Lesions, which include seborrheic keratoses and other benign growths; Actinic Keratoses, precancerous lesions that can develop into squamous cell carcinoma; Squamous Cell Carcinoma, the second most common form of skin cancer; Vascular Lesions, which include various blood vessel-related skin conditions; and Dermatofibroma, benign fibrous nodules commonly found on the lower extremities.",
        
        "The user authentication and security framework implemented in SkinCare AI ensures that sensitive health information remains protected while providing a seamless user experience. The registration process incorporates email verification using One-Time Passwords, ensuring that users provide valid email addresses and preventing unauthorized account creation. The OTP system generates secure six-digit codes that expire after ten minutes, balancing security requirements with user convenience. Password management follows industry best practices, with passwords being hashed using Django's PBKDF2 algorithm before storage, ensuring that even in the unlikely event of a database breach, user credentials remain protected.",
        
        "The analysis history and tracking functionality enables users to maintain a comprehensive record of all their skin analyses over time. This feature is particularly valuable for monitoring changes in existing lesions or tracking the appearance of new skin conditions. Each analysis record includes the original uploaded image, the classification result, confidence scores, the model used for analysis, and the timestamp of the analysis. Users can easily access their complete history through an intuitive interface that supports filtering and sorting by various criteria including date, condition type, and confidence level.",
        
        "The PDF report generation capability allows users to create professional, printable reports of their analysis results that can be shared with healthcare providers during medical consultations. These reports include all relevant information from the analysis, presented in a clear and professional format that facilitates communication between patients and their healthcare providers. The reports incorporate appropriate medical disclaimers and clearly indicate that the results are from an AI-based preliminary screening tool rather than a professional medical diagnosis.",
        
        "The analytics dashboard provides users with visual representations of their analysis data, enabling them to identify patterns and trends over time. The dashboard utilizes Chart.js to create interactive visualizations including pie charts showing the distribution of detected conditions, line graphs displaying analysis frequency over time, and bar charts comparing confidence levels across different analyses. These visualizations help users understand their skin health patterns and can be valuable for discussions with healthcare providers.",
        
        "The DermaGenie AI assistant represents an innovative feature that provides users with an intelligent conversational interface for skin health queries. Powered by the Perplexity AI platform, DermaGenie can answer questions about various skin conditions, provide general skin care guidance, explain medical terminology, and offer educational information about skin health. The chatbot maintains conversation context, allowing for natural, flowing discussions about skin health topics. Importantly, DermaGenie is programmed to consistently remind users that its responses are for educational purposes only and should not replace professional medical advice.",
        
        "The email notification system keeps users informed about important events and updates related to their account and analyses. The system utilizes the Resend API for reliable email delivery and includes features such as welcome emails for new users, notifications for first-time analyses, profile update confirmations, and OTP delivery for authentication purposes. Users have control over their notification preferences and can opt out of non-essential communications while still receiving critical security-related emails.",
        
        "The responsive design ensures that SkinCare AI provides an optimal user experience across all device types, from desktop computers to tablets and smartphones. The interface automatically adapts to different screen sizes and orientations, ensuring that all features remain accessible and usable regardless of the device being used. Special attention has been paid to touch-friendly interfaces for mobile users, with appropriately sized buttons and touch targets that facilitate easy navigation on smaller screens.",
        
        "The administrative dashboard provides system administrators with comprehensive tools for monitoring and managing the SkinCare AI platform. Administrators can view system-wide statistics including total users, total analyses performed, and recent activity trends. The dashboard also provides user management capabilities, allowing administrators to view user accounts, monitor usage patterns, and address any issues that may arise. Access to the administrative dashboard is restricted to staff members through a separate authentication flow, ensuring that sensitive system information remains protected."
    ]
    
    for para in features_paras:
        add_paragraph_text(doc, para)

    
    add_section_heading(doc, "1.3", "Project Motivation")
    
    motivation_paras = [
        "The motivation for developing SkinCare AI stems from a confluence of factors including the growing global burden of skin cancer, the limitations of current healthcare delivery systems, and the remarkable advances in artificial intelligence that have made sophisticated medical image analysis accessible to a broader audience. This section explores the various factors that motivated the development of this project and the underlying rationale for the design decisions that shaped the final system.",
        
        "Skin cancer represents a significant and growing public health challenge worldwide. According to the World Health Organization, the incidence of both melanoma and non-melanoma skin cancers has been increasing steadily over the past several decades. In the United States alone, it is estimated that one in five Americans will develop skin cancer by the age of seventy. The American Cancer Society projects that approximately 97,610 new melanomas will be diagnosed in the United States in 2023, with an estimated 7,990 deaths resulting from the disease. These statistics underscore the urgent need for effective screening and early detection mechanisms that can help identify potential skin cancers before they progress to more advanced and dangerous stages.",
        
        "The critical importance of early detection in skin cancer outcomes cannot be overstated. When melanoma is detected at its earliest stage, before it has penetrated the epidermis, the five-year survival rate exceeds ninety-nine percent. However, when melanoma is detected at later stages after it has spread to distant organs, the five-year survival rate drops dramatically to approximately twenty-seven percent. This stark difference in outcomes based on the stage of detection highlights the life-saving potential of early screening and the value of tools that can encourage individuals to seek professional evaluation when potential concerns are identified.",
        
        "Despite the clear benefits of early detection, access to dermatological expertise remains limited in many regions around the world. In the United States, there is approximately one dermatologist for every 30,000 people, with significant geographic disparities that leave many rural and underserved communities with limited access to specialized skin care. In developing countries, the situation is often more severe, with some regions having fewer than one dermatologist per million people. This shortage of dermatological expertise creates barriers to timely screening and diagnosis, potentially allowing skin cancers to progress to more advanced stages before they are identified.",
        
        "The economic burden of skin cancer treatment adds another dimension to the motivation for developing accessible screening tools. The cost of treating skin cancer in the United States is estimated at over eight billion dollars annually, with costs increasing significantly for cancers detected at later stages. Early detection not only improves patient outcomes but also reduces the overall cost of treatment, as early-stage skin cancers can often be treated with relatively simple outpatient procedures compared to the extensive treatments required for advanced cancers. By encouraging early screening and professional consultation, tools like SkinCare AI have the potential to contribute to reduced healthcare costs while improving patient outcomes.",
        
        "The rapid advancement of artificial intelligence and deep learning technologies has created unprecedented opportunities for developing sophisticated medical image analysis systems. Convolutional Neural Networks have demonstrated remarkable capabilities in image classification tasks, with some studies showing that well-trained models can achieve diagnostic accuracy comparable to or exceeding that of experienced dermatologists. The availability of large, well-curated datasets of dermoscopic images, such as the ISIC archive and the HAM10000 dataset, has facilitated the development of robust models that can generalize across diverse patient populations and imaging conditions.",
        
        "The democratization of deep learning frameworks and cloud computing resources has made it possible to develop and deploy sophisticated AI systems without requiring extensive specialized infrastructure. Frameworks such as TensorFlow and Keras provide accessible interfaces for building and training neural networks, while transfer learning techniques allow developers to leverage pre-trained models that have been trained on massive datasets. These technological advances have lowered the barriers to entry for developing medical AI applications, enabling projects like SkinCare AI to be developed within academic and research settings.",
        
        "The increasing prevalence of smartphones with high-quality cameras has created new opportunities for mobile health applications. Modern smartphones can capture images of sufficient quality for preliminary skin analysis, making it possible for users to perform initial screenings from the comfort of their homes. This accessibility is particularly valuable for individuals who may face barriers to accessing traditional healthcare services, whether due to geographic isolation, time constraints, or economic factors. By providing a web-based platform that can be accessed from any device with a camera and internet connection, SkinCare AI aims to make preliminary skin screening accessible to a broader population.",
        
        "The educational aspect of skin health awareness represents another important motivation for this project. Many individuals lack awareness of the warning signs of skin cancer and may not recognize when a skin lesion warrants professional evaluation. By providing educational information about various skin conditions alongside analysis results, SkinCare AI aims to improve skin health literacy and empower users to make informed decisions about their healthcare. The DermaGenie AI assistant further supports this educational mission by providing a conversational interface for users to learn about skin health topics.",
        
        "The responsible development of AI in healthcare contexts requires careful consideration of ethical implications and potential risks. The motivation for SkinCare AI includes a commitment to developing a system that enhances rather than replaces professional medical care. The system is designed to encourage users to seek professional evaluation when concerns are identified, rather than to provide definitive diagnoses that might discourage appropriate medical consultation. This approach reflects an understanding that AI-based screening tools are most valuable when they serve as a bridge to professional care rather than as a substitute for it."
    ]
    
    for para in motivation_paras:
        add_paragraph_text(doc, para)

    
    add_section_heading(doc, "1.4", "Scope of the Project")
    
    scope_paras = [
        "The scope of the SkinCare AI project encompasses the design, development, and deployment of a comprehensive web-based skin lesion classification system that leverages deep learning technologies to provide preliminary assessments of skin conditions. This section delineates the boundaries of the project, clarifying what is included within its scope and what falls outside the intended functionality of the system.",
        
        "The primary scope of SkinCare AI includes the development of a dual-model deep learning system capable of classifying dermoscopic images into eight distinct categories of skin conditions. The system is designed to process images uploaded by users through a web interface, apply appropriate preprocessing transformations, and generate classification predictions along with confidence scores. The classification categories include Melanocytic Nevi, Melanoma, Basal Cell Carcinoma, Benign Keratosis-like Lesions, Actinic Keratoses, Squamous Cell Carcinoma, Vascular Lesions, and Dermatofibroma. The system provides educational information about each condition and recommends appropriate next steps based on the classification results.",
        
        "The web application development scope includes the creation of a complete user-facing interface built using the Django framework. This encompasses user registration and authentication functionality with email verification, user profile management, image upload and analysis interfaces, analysis history tracking and management, PDF report generation, analytics dashboard with interactive visualizations, and the DermaGenie AI chatbot interface. The application is designed to be responsive and accessible across desktop, tablet, and mobile devices.",
        
        "The machine learning component scope includes the training and optimization of two distinct neural network models. The primary model utilizes the EfficientNetB0 architecture with transfer learning from ImageNet weights, fine-tuned on the ISIC 2019 dataset. The secondary model is a custom Convolutional Neural Network trained on a modified HAM10000 dataset. The scope includes the development of an intelligent model selection mechanism that automatically chooses the most appropriate model based on confidence thresholds and user preferences.",
        
        "The administrative functionality scope includes the development of a staff-only dashboard for system monitoring and management. This dashboard provides system-wide statistics, user management capabilities, and analysis monitoring tools. The administrative interface is protected by a separate authentication flow to ensure that sensitive system information remains accessible only to authorized personnel.",
        
        "The email notification system scope includes the integration of the Resend API for reliable email delivery. The system supports various notification types including welcome emails, analysis notifications, profile update confirmations, and OTP delivery for authentication purposes. Users have control over their notification preferences through their profile settings.",
        
        "The security implementation scope includes the development of comprehensive security measures to protect user data and ensure system integrity. This encompasses password hashing, session management, CSRF protection, input validation, secure file upload handling, and environment variable management for sensitive configuration data.",
        
        "Several important limitations define what falls outside the scope of SkinCare AI. The system is explicitly not intended to provide medical diagnoses or treatment recommendations. All results are presented as preliminary assessments for educational purposes only, with clear disclaimers emphasizing the need for professional medical evaluation. The system does not replace the expertise of qualified dermatologists or other healthcare providers, and users are consistently encouraged to seek professional consultation for any skin concerns.",
        
        "The scope does not include the analysis of non-dermoscopic images or images of conditions outside the eight supported categories. While the system can process various image types, it has been trained specifically on dermoscopic images and may not provide accurate results for other image types. The system does not support real-time video analysis or webcam-based screening, although these features are identified as potential future enhancements.",
        
        "The current scope is limited to a web-based application and does not include native mobile applications for iOS or Android platforms. While the web application is designed to be mobile-responsive and accessible from mobile browsers, dedicated mobile applications with features such as offline analysis or camera integration are outside the current scope.",
        
        "The scope does not include integration with electronic health record systems or telemedicine platforms. While the PDF report generation feature facilitates sharing of results with healthcare providers, direct integration with clinical systems is not included in the current implementation.",
        
        "The geographic scope of the system is not limited to any specific region, although the interface is currently available only in English. Internationalization and multi-language support are identified as potential future enhancements but are not included in the current scope.",
        
        "The scope includes deployment on local development servers and provides guidance for production deployment, but does not include the actual deployment to production cloud infrastructure or the ongoing maintenance and support of a production system."
    ]
    
    for para in scope_paras:
        add_paragraph_text(doc, para)

    
    add_section_heading(doc, "1.5", "Organization of the Report")
    
    org_paras = [
        "This project report is organized into ten chapters, each addressing a specific aspect of the SkinCare AI system development. The structure follows a logical progression from introduction and analysis through design, implementation, testing, and conclusion, providing a comprehensive documentation of the entire project lifecycle.",
        
        "Chapter One, Introduction, provides an overview of the SkinCare AI project, including the context and background of skin cancer detection, the salient features of the developed system, the motivation behind the project, the scope and limitations, and the organization of this report. This chapter establishes the foundation for understanding the subsequent technical discussions.",
        
        "Chapter Two, System Study and Analysis, presents a detailed analysis of the problem domain and the requirements for the proposed solution. This chapter includes the formal problem statement, an examination of existing systems and their limitations, a description of the proposed system and its advantages, and a comprehensive feasibility analysis covering technical, economic, and operational aspects.",
        
        "Chapter Three, Development Environment, documents the hardware and software requirements for developing and running the SkinCare AI system. This chapter provides detailed descriptions of the key technologies used in the project, including Python, Google Colab, Django, TensorFlow, and the database technology employed.",
        
        "Chapter Four, System Design, presents the architectural and design decisions that shaped the SkinCare AI system. This chapter includes detailed module descriptions for all major system components, the methodology employed in development, input and output design specifications, data flow diagrams at multiple levels, the system architecture diagram, and the database design including entity-relationship diagrams and table descriptions.",
        
        "Chapter Five, System Implementation, documents the actual implementation of the SkinCare AI system. This chapter covers frontend implementation details, backend implementation specifics, machine learning model implementation, the model training process, model evaluation metrics, and security implementation measures.",
        
        "Chapter Six, System Testing, describes the testing strategies and procedures employed to ensure the quality and reliability of the SkinCare AI system. This chapter covers the overall testing strategy, unit testing procedures, integration testing approaches, validation testing methods, test case design principles, and sample test cases with expected and actual results.",
        
        "Chapter Seven, Results and Discussion, presents the experimental results obtained from the SkinCare AI system and provides analysis and interpretation of these results. This chapter includes experimental results from model training and evaluation, performance analysis of the system, accuracy analysis across different conditions and scenarios, and comparison with existing systems in the literature.",
        
        "Chapter Eight, Conclusion and Future Enhancement, summarizes the achievements of the SkinCare AI project and identifies opportunities for future development. This chapter provides a comprehensive conclusion of the work accomplished and outlines potential enhancements that could extend the functionality and impact of the system.",
        
        "Chapter Nine, Bibliography and References, provides a comprehensive list of all sources cited in this report, including academic papers, technical documentation, and online resources that informed the development of SkinCare AI.",
        
        "Chapter Ten, Appendices, contains supplementary materials that support the main body of the report. Appendix A presents output screens showing the user interface of the SkinCare AI system. Appendix B provides sample code snippets illustrating key implementation details. Appendix C contains a user manual providing guidance on how to use the SkinCare AI system.",
        
        "The organization of this report is designed to provide readers with a clear and logical progression through the project, from initial concept and analysis through design, implementation, and evaluation. Each chapter builds upon the preceding chapters, creating a comprehensive documentation of the SkinCare AI project that can serve as both a technical reference and an academic record of the work accomplished."
    ]
    
    for para in org_paras:
        add_paragraph_text(doc, para)
    
    doc.add_page_break()


def add_remaining_chapters(doc):
    """Add remaining chapters with key content"""
    
    # Chapter 2
    add_chapter_heading(doc, "2", "SYSTEM STUDY AND ANALYSIS")
    
    add_section_heading(doc, "2.1", "Problem Statement")
    add_paragraph_text(doc, "The global healthcare landscape faces a significant challenge in the early detection and diagnosis of skin cancer, a disease that affects millions of individuals worldwide and claims thousands of lives annually. Despite advances in medical technology and increased awareness of skin cancer risks, substantial barriers continue to impede timely diagnosis and treatment, particularly in underserved communities and regions with limited access to dermatological expertise.")
    add_paragraph_text(doc, "The primary problem can be stated as follows: There exists a critical need for accessible, accurate, and user-friendly preliminary screening tools that can assist individuals in identifying potentially concerning skin lesions and encourage timely professional medical evaluation. The current healthcare infrastructure is insufficient to provide universal access to dermatological screening, resulting in delayed diagnoses, poorer patient outcomes, and increased healthcare costs associated with treating advanced-stage skin cancers.")
    add_paragraph_text(doc, "The problem manifests across multiple dimensions including the epidemiological burden of skin cancer, geographic and socioeconomic disparities in access to dermatological care, the time-sensitive nature of skin cancer diagnosis, limitations of self-examination, and the economic impact of delayed diagnosis and treatment.")
    
    add_section_heading(doc, "2.2", "Existing System")
    add_paragraph_text(doc, "Traditional clinical diagnosis of skin lesions relies primarily on visual examination by trained healthcare providers, typically dermatologists or primary care physicians with dermatological training. Dermoscopy represents an advancement over naked-eye examination, using a handheld device to examine skin lesions with magnification and specialized lighting. Histopathological examination remains the gold standard for definitive diagnosis.")
    add_paragraph_text(doc, "Several commercial and research-based AI systems for skin lesion analysis have emerged in recent years. Teledermatology services have emerged as a means of extending dermatological expertise to underserved areas. Consumer-facing mobile applications for skin analysis have proliferated with varying levels of sophistication and accuracy.")
    
    add_subsection_heading(doc, "2.2.1", "Drawbacks of Existing System")
    add_paragraph_text(doc, "The limited accessibility of professional dermatological care represents a fundamental drawback of traditional clinical approaches. The subjective nature of visual examination introduces variability in diagnostic accuracy. Existing AI-based systems often suffer from limited transparency and explainability. Many existing consumer applications lack comprehensive features beyond basic image classification. The absence of proper medical disclaimers and responsible AI practices raises concerns about potential misuse.")
    
    add_section_heading(doc, "2.3", "Proposed System")
    add_paragraph_text(doc, "SkinCare AI represents a comprehensive solution designed to address the limitations of existing skin cancer detection systems while providing an accessible, accurate, and user-friendly platform for preliminary skin lesion screening. The core of the proposed system is a dual-model deep learning architecture that combines the strengths of two distinct neural network models to provide robust and accurate skin lesion classification.")
    add_paragraph_text(doc, "The intelligent model selection mechanism represents a key innovation of the proposed system. Rather than relying on a single model for all predictions, the system automatically evaluates the confidence level of predictions from the primary model and determines whether to utilize the secondary model for potentially improved accuracy.")
    
    add_subsection_heading(doc, "2.3.1", "Advantages of Proposed System")
    add_paragraph_text(doc, "The dual-model architecture provides improved accuracy and robustness compared to single-model systems. The comprehensive feature set extends well beyond basic image classification. The emphasis on responsible AI practices ensures that the system enhances rather than replaces professional medical care. The user-friendly interface makes advanced AI technology accessible to users without technical expertise.")
    
    add_section_heading(doc, "2.4", "Feasibility Analysis")
    
    add_subsection_heading(doc, "2.4.1", "Technical Feasibility")
    add_paragraph_text(doc, "The deep learning component of the system relies on well-established technologies that have been extensively validated in both research and production environments. TensorFlow is a mature, open-source platform developed by Google. The web application component utilizes Django, a mature and well-documented web framework. Based on this analysis, the SkinCare AI system is determined to be technically feasible.")
    
    add_subsection_heading(doc, "2.4.2", "Economic Feasibility")
    add_paragraph_text(doc, "The development costs for the SkinCare AI system are primarily associated with personnel time, as the project utilizes open-source technologies that do not require licensing fees. The machine learning model training can be conducted using Google Colab, which provides free access to GPU-accelerated computing resources. The cost-benefit analysis indicates that the SkinCare AI system is economically feasible.")
    
    add_subsection_heading(doc, "2.4.3", "Operational Feasibility")
    add_paragraph_text(doc, "The user interface has been designed following established usability principles to ensure that users can effectively navigate the application. The training requirements for using the system are minimal. The system has been designed to complement rather than replace existing healthcare workflows. Based on this analysis, the SkinCare AI system is determined to be operationally feasible.")
    
    doc.add_page_break()

    
    # Chapter 3
    add_chapter_heading(doc, "3", "DEVELOPMENT ENVIRONMENT")
    
    add_section_heading(doc, "3.1", "Hardware Requirements")
    add_paragraph_text(doc, "The hardware requirements for the SkinCare AI system vary depending on the intended use case, ranging from development and testing environments to production deployment scenarios. The minimum hardware requirements for development include a processor with at least four cores, such as an Intel Core i5 or AMD Ryzen 5, a minimum of eight gigabytes of RAM, and at least ten gigabytes of free disk space.")
    
    # Add hardware requirements table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Component', 'Minimum', 'Recommended']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        for para in table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    
    hw_data = [
        ('Processor', 'Intel Core i5 / AMD Ryzen 5', 'Intel Core i7 / AMD Ryzen 7'),
        ('RAM', '8 GB', '16 GB'),
        ('Storage', '10 GB free space', '20 GB SSD'),
        ('GPU', 'Not required', 'NVIDIA GTX 1060 or better'),
        ('Display', '1366 x 768', '1920 x 1080'),
    ]
    for i, (comp, min_req, rec_req) in enumerate(hw_data, 1):
        table.rows[i].cells[0].text = comp
        table.rows[i].cells[1].text = min_req
        table.rows[i].cells[2].text = rec_req
        for cell in table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Table 3.1: Hardware Requirements")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    
    doc.add_paragraph()
    
    add_section_heading(doc, "3.2", "Software Requirements")
    add_paragraph_text(doc, "The software requirements for the SkinCare AI system encompass the operating system, programming languages, frameworks, libraries, and external services needed to develop, test, and deploy the application. Python version 3.10 or later is required. The Django framework version 4.2.1 provides the foundation for the web application. TensorFlow version 2.13.0 serves as the deep learning framework.")
    
    add_section_heading(doc, "3.3", "Software Description")
    
    add_subsection_heading(doc, "3.3.1", "About Python")
    add_paragraph_text(doc, "Python is a high-level, interpreted programming language that has become one of the most popular languages for software development, particularly in the fields of data science, machine learning, and web development. The design philosophy of Python emphasizes code readability through the use of significant whitespace and a clean, expressive syntax. Python's role in the SkinCare AI project is central, serving as the primary programming language for both the web application and the machine learning components.")
    
    add_subsection_heading(doc, "3.3.2", "About Google Colab")
    add_paragraph_text(doc, "Google Colaboratory, commonly known as Google Colab, is a cloud-based interactive computing environment that allows users to write and execute Python code through a web browser. The primary advantage of Google Colab for machine learning projects is the provision of free GPU and TPU resources. In the SkinCare AI project, Google Colab was used for training the deep learning models.")
    
    add_subsection_heading(doc, "3.3.3", "About Django")
    add_paragraph_text(doc, "Django is a high-level Python web framework that enables rapid development of secure and maintainable web applications. Django follows the Model-View-Template architectural pattern. The Object-Relational Mapping provided by Django abstracts database operations. Django's built-in authentication system provides comprehensive user management functionality. In the SkinCare AI project, Django serves as the foundation for the web application.")
    
    add_subsection_heading(doc, "3.3.4", "About TensorFlow")
    add_paragraph_text(doc, "TensorFlow is an open-source machine learning platform developed by the Google Brain team. The core abstraction in TensorFlow is the computational graph. Keras, which is integrated into TensorFlow as its high-level API, provides an intuitive interface for building and training neural networks. In the SkinCare AI project, TensorFlow and Keras are used for loading and executing the trained neural network models.")
    
    add_subsection_heading(doc, "3.3.5", "About Database Technology")
    add_paragraph_text(doc, "The system utilizes SQLite as the default database for development and testing, with support for migration to PostgreSQL for production deployments. SQLite is a self-contained, serverless, zero-configuration database engine. Django's Object-Relational Mapping provides a consistent interface for database operations regardless of the underlying database backend.")
    
    doc.add_page_break()

    
    # Chapter 4
    add_chapter_heading(doc, "4", "SYSTEM DESIGN")
    
    add_section_heading(doc, "4.1", "Module Description")
    add_paragraph_text(doc, "The SkinCare AI system is organized into eight primary modules, each responsible for a specific aspect of the application's functionality. This modular architecture promotes separation of concerns, facilitates maintenance and testing, and enables independent development of different system components.")
    
    add_paragraph_text(doc, "The User Authentication Module is responsible for managing user identity and access control. The Image Analysis Module is the core component responsible for processing uploaded images and generating classification predictions. The User Profile Module manages user profile information. The Analytics Dashboard Module provides data visualization and trend analysis. The History and Comparison Module enables users to access their complete analysis history. The Admin Module provides system administration and monitoring. The DermaGenie AI Assistant Module provides an AI-powered chatbot for skin health queries. The Email Notification Module handles automated email communications.")
    
    add_section_heading(doc, "4.2", "Methodology")
    add_paragraph_text(doc, "The development of the SkinCare AI system followed an iterative and incremental methodology that combined elements of agile development practices with structured analysis and design phases. The project began with a requirements gathering phase, followed by system analysis, design, implementation, and testing phases.")
    
    add_section_heading(doc, "4.3", "Input Design")
    add_paragraph_text(doc, "Input design encompasses the specification of all data inputs to the SkinCare AI system, including user-provided data through forms and file uploads, as well as the preprocessing of image data for model inference. The user registration form collects essential information for account creation. The image upload interface is the primary input mechanism for the core analysis functionality.")
    
    add_section_heading(doc, "4.4", "Output Design")
    add_paragraph_text(doc, "Output design specifies how the SkinCare AI system presents information to users. The analysis results presentation is the primary output of the system. The results presentation follows a legally-compliant format that clearly communicates the preliminary nature of the assessment. The PDF report output provides a professional, printable document summarizing analysis results.")
    
    add_section_heading(doc, "4.5", "Data Flow Diagram")
    add_paragraph_text(doc, "Data Flow Diagrams provide a graphical representation of how data moves through the SkinCare AI system, illustrating the processes that transform data, the data stores that hold information, and the external entities that interact with the system.")
    
    add_diagram_placeholder(doc, "4.2", "Data Flow Diagram Level 0 (Context Diagram)")
    add_paragraph_text(doc, "The Context Diagram provides a high-level view of the SkinCare AI system, showing it as a single process that interacts with external entities including Users, Email Service (Resend), Perplexity AI, Admin Users, and ML Models.")
    
    add_diagram_placeholder(doc, "4.3", "Data Flow Diagram Level 1")
    add_paragraph_text(doc, "The Level 1 DFD decomposes the SkinCare AI system into its major processes: User Authentication, Image Analysis, ML Model Inference, History Management, Profile Management, and Admin Dashboard.")
    
    add_section_heading(doc, "4.6", "Architecture Diagram")
    add_diagram_placeholder(doc, "4.1", "System Architecture Diagram")
    add_paragraph_text(doc, "The System Architecture Diagram shows the layered architecture of SkinCare AI including the Presentation Layer (client browsers, Django templates), Application Layer (Django framework, business logic), Machine Learning Layer (TensorFlow/Keras, EfficientNetB0, Custom CNN), Data Layer (SQLite database, file storage), and External Services (Resend API, Perplexity AI).")
    
    add_section_heading(doc, "4.7", "Database Design")
    
    add_subsection_heading(doc, "4.7.1", "ER Diagram")
    add_diagram_placeholder(doc, "4.4", "Entity Relationship Diagram")
    add_paragraph_text(doc, "The Entity-Relationship Diagram illustrates the logical structure of the database, showing the entities (User, UserProfile, UserPredictModel, EmailOTP, PasswordResetOTP, ChatConversation), their attributes, and the relationships between them.")
    
    add_subsection_heading(doc, "4.7.2", "Table Description")
    add_paragraph_text(doc, "The database schema includes tables for user accounts (auth_user), user profiles (APP_userprofile), analysis records (APP_userpredictmodel), email verification tokens (APP_emailotp), password reset tokens (APP_passwordresetotp), and chat conversation history (APP_chatconversation).")
    
    doc.add_page_break()

    
    # Chapter 5
    add_chapter_heading(doc, "5", "SYSTEM IMPLEMENTATION")
    
    add_section_heading(doc, "5.1", "Frontend Implementation")
    add_paragraph_text(doc, "The frontend implementation of SkinCare AI encompasses all user-facing components of the application, including HTML templates, CSS stylesheets, and JavaScript functionality. The template architecture follows Django's template inheritance system. The visual design follows a dark, futuristic theme that conveys professionalism and technological sophistication. The responsive design implementation ensures that the application functions well on devices ranging from large desktop monitors to small smartphone screens.")
    
    add_section_heading(doc, "5.2", "Backend Implementation")
    add_paragraph_text(doc, "The backend implementation of SkinCare AI is built on the Django framework, providing a robust foundation for handling HTTP requests, managing database operations, and implementing business logic. The project structure follows Django's recommended organization. The view implementation handles request processing and response generation. The utility modules encapsulate reusable functionality.")
    
    add_section_heading(doc, "5.3", "Machine Learning Model Implementation")
    add_paragraph_text(doc, "The machine learning model implementation encompasses the integration of trained neural network models into the SkinCare AI web application. The model loading process occurs during application startup. The image preprocessing pipeline transforms uploaded images into the format expected by the neural network models. The intelligent model selection mechanism implements the logic for choosing between the primary and secondary models.")
    
    add_section_heading(doc, "5.4", "Model Training Process")
    add_paragraph_text(doc, "The model training process was conducted using Google Colab, which provided free access to GPU-accelerated computing resources. The data preparation phase involved downloading and preprocessing the training datasets. Data augmentation techniques included random rotation, flipping, zoom, and brightness adjustment. The EfficientNetB0 model training used transfer learning from ImageNet weights. The custom CNN model architecture consists of three convolutional blocks followed by fully connected layers.")
    
    add_section_heading(doc, "5.5", "Model Evaluation Metrics")
    add_paragraph_text(doc, "The model evaluation phase assessed the performance of the trained models using various metrics. The overall accuracy metric measures the proportion of correct predictions. The EfficientNetB0 model achieved a test accuracy of 71.32 percent on the ISIC 2019 dataset. The custom CNN model achieved a test accuracy of 94.1 percent on the HAM10000 subset. Additional metrics include precision, recall, F1 score, and confusion matrix analysis.")
    
    add_section_heading(doc, "5.6", "Security Implementation")
    add_paragraph_text(doc, "The security implementation encompasses multiple layers of protection. Password security is implemented using Django's built-in password hashing system with PBKDF2 algorithm. Session management uses Django's session framework with secure cookie settings. Cross-Site Request Forgery protection is enabled by default. SQL injection prevention is achieved through the use of Django's ORM. File upload security includes validation of uploaded files.")
    
    doc.add_page_break()
    
    # Chapter 6
    add_chapter_heading(doc, "6", "SYSTEM TESTING")
    
    add_section_heading(doc, "6.1", "Testing Strategy")
    add_paragraph_text(doc, "The testing strategy for SkinCare AI encompasses multiple levels of testing to ensure that the system functions correctly, meets requirements, and provides a satisfactory user experience. The testing approach follows the testing pyramid model, with a broad base of unit tests, a middle layer of integration tests, and end-to-end tests at the top.")
    
    add_section_heading(doc, "6.2", "Unit Testing")
    add_paragraph_text(doc, "Unit testing focuses on testing individual components in isolation. The model unit tests verify that the machine learning models produce expected outputs. The view unit tests verify that view functions handle requests correctly. The form unit tests verify that form validation works correctly.")
    
    add_section_heading(doc, "6.3", "Integration Testing")
    add_paragraph_text(doc, "Integration testing verifies that different components of the system work together correctly. The authentication integration tests verify the complete authentication workflow. The analysis integration tests verify the complete image analysis workflow. The profile integration tests verify profile management functionality.")
    
    add_section_heading(doc, "6.4", "Validation Testing")
    add_paragraph_text(doc, "Validation testing verifies that the system meets the specified requirements. The functional requirements validation verifies that all specified features are implemented. The user interface validation verifies usability requirements. The performance validation verifies response time requirements. The security validation verifies security requirements.")
    
    add_section_heading(doc, "6.5", "Test Case Design")
    add_paragraph_text(doc, "Test case design follows established principles including equivalence partitioning, boundary value analysis, decision table technique, state transition technique, and error guessing technique.")
    
    add_section_heading(doc, "6.6", "Sample Test Cases")
    add_paragraph_text(doc, "Sample test cases cover user authentication (registration, login, OTP verification, password reset), image analysis (valid image upload, invalid file type, model selection), profile management (view profile, update profile, upload picture), and system integration (complete registration flow, analysis to history flow, PDF generation).")
    
    doc.add_page_break()

    
    # Chapter 7
    add_chapter_heading(doc, "7", "RESULTS AND DISCUSSION")
    
    add_section_heading(doc, "7.1", "Experimental Results")
    add_paragraph_text(doc, "The experimental results of the SkinCare AI system demonstrate the effectiveness of the dual-model approach for skin lesion classification. The EfficientNetB0 model training was conducted on the ISIC 2019 dataset comprising 25,331 dermoscopic images. The final test accuracy for the EfficientNetB0 model was 71.32 percent. The custom CNN model training was conducted on the modified HAM10000 subset comprising 5,906 images. The final test accuracy for the custom CNN model was 94.1 percent.")
    
    add_section_heading(doc, "7.2", "Performance Analysis")
    add_paragraph_text(doc, "The performance analysis examines the computational efficiency and resource utilization of the system. The EfficientNetB0 model inference time averaged 850 milliseconds on CPU. The custom CNN model inference time averaged 120 milliseconds on CPU. The total application memory footprint is approximately 500 megabytes. The system maintained acceptable response times with up to 50 concurrent users.")
    
    add_section_heading(doc, "7.3", "Accuracy Analysis")
    add_paragraph_text(doc, "The accuracy analysis provides detailed examination of model performance across different conditions. The per-class accuracy analysis for EfficientNetB0 reveals variation across conditions, with Melanocytic Nevi achieving 82 percent and Melanoma achieving 68 percent. The custom CNN shows more uniform performance with Melanocytic Nevi at 96 percent and Melanoma at 92 percent. Higher confidence predictions tend to be more accurate.")
    
    add_section_heading(doc, "7.4", "Comparison with Existing Systems")
    add_paragraph_text(doc, "The comparison with existing systems contextualizes the performance of SkinCare AI. The EfficientNetB0 model's 71.32 percent accuracy on ISIC 2019 is comparable to results reported in recent literature (65-80 percent range). The custom CNN's 94.1 percent accuracy exceeds many reported results on similar datasets. The comprehensive feature set distinguishes SkinCare AI from applications that provide only basic classification.")
    
    doc.add_page_break()
    
    # Chapter 8
    add_chapter_heading(doc, "8", "CONCLUSION AND FUTURE ENHANCEMENT")
    
    add_section_heading(doc, "8.1", "Conclusion")
    add_paragraph_text(doc, "The SkinCare AI project has successfully achieved its primary objective of developing a comprehensive, accessible, and responsible AI-powered skin lesion classification system. The dual-model architecture represents a significant technical achievement, combining EfficientNetB0 trained on ISIC 2019 with a custom CNN trained on HAM10000 subset.")
    add_paragraph_text(doc, "The web application provides a user-friendly interface that makes advanced AI technology accessible to users without technical expertise. The comprehensive feature set includes user authentication, analysis history tracking, PDF report generation, analytics dashboard, and AI chatbot assistance.")
    add_paragraph_text(doc, "The emphasis on responsible AI practices ensures that SkinCare AI serves as a complement to professional medical care. The experimental results demonstrate competitive accuracy on standard benchmark datasets. The project has demonstrated the feasibility of developing sophisticated medical AI applications using open-source technologies.")
    
    add_section_heading(doc, "8.2", "Future Enhancements")
    add_paragraph_text(doc, "Future enhancements include: development of native mobile applications for iOS and Android; model improvements through training on larger datasets and ensemble methods; real-time analysis using webcam or smartphone camera feeds; telemedicine integration for direct connection with dermatologists; multi-language support for international accessibility; advanced analytics with predictive capabilities; API development for third-party integration; and research platform capabilities for contributing to skin cancer research.")
    
    doc.add_page_break()

    
    # Chapter 9 - References
    add_chapter_heading(doc, "9", "BIBLIOGRAPHY AND REFERENCES")
    
    references = [
        "1. Codella, N. C. F., et al. (2018). Skin lesion analysis toward melanoma detection: A challenge at the 2017 ISBI. IEEE 15th International Symposium on Biomedical Imaging.",
        "2. Esteva, A., et al. (2017). Dermatologist-level classification of skin cancer with deep neural networks. Nature, 542(7639), 115-118.",
        "3. Haenssle, H. A., et al. (2018). Man against machine: diagnostic performance of a deep learning CNN for dermoscopic melanoma recognition. Annals of Oncology, 29(8), 1836-1842.",
        "4. Tan, M., & Le, Q. (2019). EfficientNet: Rethinking model scaling for convolutional neural networks. ICML.",
        "5. Tschandl, P., Rosendahl, C., & Kittler, H. (2018). The HAM10000 dataset. Scientific Data, 5(1), 1-9.",
        "6. World Health Organization. (2023). Skin cancers. https://www.who.int/",
        "7. American Cancer Society. (2023). Cancer Facts & Figures 2023.",
        "8. Skin Cancer Foundation. (2023). Skin Cancer Facts & Statistics.",
        "9. Django Software Foundation. (2023). Django Documentation. https://docs.djangoproject.com/",
        "10. TensorFlow. (2023). TensorFlow Documentation. https://www.tensorflow.org/",
        "11. Keras. (2023). Keras Documentation. https://keras.io/",
        "12. International Skin Imaging Collaboration. (2023). ISIC Archive. https://www.isic-archive.com/",
        "13. Chollet, F. (2017). Deep Learning with Python. Manning Publications.",
        "14. Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.",
        "15. LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep learning. Nature, 521(7553), 436-444.",
    ]
    
    for ref in references:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    
    doc.add_page_break()
    
    # Chapter 10 - Appendices
    add_chapter_heading(doc, "10", "APPENDICES")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("APPENDIX A: OUTPUT SCREENS")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    
    screens = [
        ("A.1", "Landing Page"),
        ("A.2", "User Registration Page"),
        ("A.3", "Login Page"),
        ("A.4", "Home Dashboard"),
        ("A.5", "Image Analysis Page"),
        ("A.6", "Analysis Results Page"),
        ("A.7", "User Profile Page"),
        ("A.8", "Analytics Dashboard"),
    ]
    
    for fig_num, title in screens:
        add_diagram_placeholder(doc, fig_num, title)
    
    doc.add_page_break()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("APPENDIX B: SAMPLE CODE SNIPPETS")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    add_paragraph_text(doc, "This appendix contains sample code snippets illustrating key implementation details of the SkinCare AI system, including model loading and prediction, user registration view, OTP generation and verification, email notification utility, and Django model definitions.")
    add_paragraph_text(doc, "[Code snippets to be inserted here - See PROJECT_REPORT.md Appendix B for complete code listings]")
    
    doc.add_page_break()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("APPENDIX C: USER MANUAL")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    
    manual_sections = [
        "Getting Started: Navigate to the application URL in your web browser.",
        "Creating an Account: Click Register, enter username, email, password, verify email with OTP.",
        "Logging In: Enter username/email and password, click Login.",
        "Performing a Skin Analysis: Upload image, select model preference, click Analyze, review results.",
        "Viewing Analysis History: Navigate to History page, browse past analyses, use filters.",
        "Generating PDF Reports: Navigate to analysis result, click Generate PDF.",
        "Using Analytics Dashboard: View condition distribution charts and analysis trends.",
        "Managing Your Profile: Update personal information and notification preferences.",
        "Using DermaGenie AI Assistant: Type questions about skin health, receive AI responses.",
        "Resetting Your Password: Click Forgot Password, enter email, verify OTP, set new password.",
        "Important Reminders: Results are for educational purposes only. Always consult a qualified dermatologist."
    ]
    
    for section in manual_sections:
        p = doc.add_paragraph()
        run = p.add_run("• " + section)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("--- END OF PROJECT REPORT ---")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True


def main():
    """Main function to generate the DOCX file"""
    print("Creating PROJECT_REPORT_FINAL.docx...")
    
    # Create document
    doc = Document()
    
    # Set page margins
    set_page_margins(doc)
    
    # Add all sections
    print("Adding title page...")
    add_title_page(doc)
    
    print("Adding certificate...")
    add_certificate(doc)
    
    print("Adding declaration...")
    add_declaration(doc)
    
    print("Adding acknowledgement...")
    add_acknowledgement(doc)
    
    print("Adding abstract...")
    add_abstract(doc)
    
    print("Adding table of contents...")
    add_toc(doc)
    
    print("Adding list of tables...")
    add_list_of_tables(doc)
    
    print("Adding list of figures...")
    add_list_of_figures(doc)
    
    print("Adding Chapter 1: Introduction...")
    add_chapter1(doc)
    
    print("Adding remaining chapters...")
    add_remaining_chapters(doc)
    
    # Save document
    output_path = "docs/PROJECT_REPORT_FINAL.docx"
    doc.save(output_path)
    print(f"\nDocument saved successfully: {output_path}")
    print("The document is now ready for final review and printing.")

if __name__ == "__main__":
    main()
